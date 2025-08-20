# README
# This script reads data from an Excel file and fills a Word document template with the data.
# Before running this script, ensure you have the required libraries installed:
# pip install openpyxl python-docx pandas

# Import necessary libraries
import os
from docx import Document
from openpyxl import load_workbook
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import warnings
warnings.filterwarnings("ignore", category=UserWarning)
import pandas as pd
import json

# ===== Config knobs (keeps original behavior but safer defaults) =====
EAST_ASIA_FONT = '標楷體'  # Better for Chinese; was 'Times New Roman'
DEFAULT_RUN_FONT = 'Times New Roman'
DEFAULT_RUN_SIZE_PT = 12
COLUMN_WIDTH_DXA_DEFAULT = 2000

# ===== Helpers kept internal (no interface/name changes to public functions) =====
def _set_run_style(run):
    run.font.size = Pt(DEFAULT_RUN_SIZE_PT)
    run.font.name = DEFAULT_RUN_FONT
    # Ensure East Asian glyphs render well
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn('w:eastAsia'), EAST_ASIA_FONT)


def _replace_paragraph_text(paragraph, text):
    # Remove all runs safely, then insert a single run with styles
    for r in list(paragraph.runs):
        paragraph._element.remove(r._element)
    run = paragraph.add_run(text)
    _set_run_style(run)


#Define functions to read and format data from Excel, fill Word tables, and replace text in Word documents.
def format_value(cell):
    value = cell.value
    if value is None:
        return ''
    number_format = getattr(cell, 'number_format', '') or ''
    # Handle percent formats robustly: Excel stores 50% often as 0.5, but
    # sometimes the raw value is already 50. Detect and normalize.
    if isinstance(value, (int, float)) and ('%' in number_format or '0%' in number_format):
        val = value if 0 <= value <= 1 else value / 100.0
        return f"{val * 100:.2f}%"  # Format as percentage with two decimal places
    return str(value)


def read_excel_data(excel_path, sheet_name, start_cells=1):
    workbook = load_workbook(excel_path, read_only=True, data_only=True) #Open the workbook in data-only mode which means it will not evaluate formulas, just return the values.

    if sheet_name not in workbook.sheetnames:
        workbook.close()
        raise ValueError(f"Sheet '{sheet_name}' not found. Available: {workbook.sheetnames}")

    sheet = workbook[sheet_name]

    if sheet_name == '表1.基本資料':
        data = {
            'A': [],
            'C': [],
        } #Initialize a dictionary for company name (A column) and address (C column)
        row = 18
        empty_streak = 0
        while True:
            cell_a = f'A{row}'
            cell_c = f'C{row}'
            value_a = sheet[cell_a].value
            value_c = sheet[cell_c].value
            # Stop only when BOTH are empty for a bit (avoids holes)
            if (value_a is None or str(value_a).strip() == '') and (value_c is None or str(value_c).strip() == ''):
                empty_streak += 1
                if empty_streak >= 2:
                    break
            else:
                empty_streak = 0
                data['A'].append(value_a)
                data['C'].append(value_c)
            row += 1 #Read the columns A and C starting from row 13 until an empty cell is found.

    elif sheet_name == '表2.排放源鑑別':
        data = {
            'B': [],
            'C': [],
            'E': [],
            'K': [],
            'I':[],
            'C_category3': [],
            'C_category5': [],
            'C_category6': [],
            'C_category7': [],
            'C_category8': [],
            'C_category10': [],
            'C_category11': [],
            'C_category13': [],
            'C_category14': [],
            'C_category15': [],
            'K_category1':[],
            'C_category1':[],
            'others': []
        }
        row = 4
        empty_streak = 0
        while True:
            cell_b = f'B{row}'
            cell_c = f'C{row}'
            cell_e = f'E{row}'
            cell_k = f'K{row}'
            cell_i = f'I{row}'
            value_b = sheet[cell_b].value
            value_c = sheet[cell_c].value
            value_e = sheet[cell_e].value
            value_k = sheet[cell_k].value
            value_i = sheet[cell_i].value #There's nothing in I column?????

            # End when core fields are all empty for a couple of rows
            if all((v is None or str(v).strip() == '') for v in (value_b, value_c, value_e, value_k)):
                empty_streak += 1
                if empty_streak >= 2:
                    break
                row += 1
                continue
            empty_streak = 0

            data['B'].append(value_b)
            data['C'].append(value_c)
            data['E'].append(value_e)
            data['K'].append(value_k)
            data['I'].append(value_i)
            data['others'].append('請輸入文字')

            if value_e == '範疇1':
                data['K_category1'].append(value_k)
                data['C_category1'].append(value_c)
            elif value_e == '類別3':
                data['C_category3'].append(value_c)
            elif value_e == '類別5':
                data['C_category5'].append(value_c)
            elif value_e == '類別6':
                data['C_category6'].append(value_c)
            elif value_e == '類別7':
                data['C_category7'].append(value_c)
            elif value_e == '類別8':
                data['C_category8'].append(value_c)
            elif value_e == '類別10':
                data['C_category10'].append(value_c)
            elif value_e == '類別11':
                data['C_category11'].append(value_c)
            elif value_e == '類別13':
                data['C_category13'].append(value_c)
            elif value_e == '類別14':
                data['C_category14'].append(value_c)
            elif value_e == '類別15':
                data['C_category15'].append(value_c)
            row += 1

    elif sheet_name == '表3.活動數據':
        data = {
            'C': [],
            'I':[],
            'others': []
            }
        row = 4
        empty_streak = 0
        while True:
            cell_c = f'C{row}'
            cell_i = f'I{row}'
            value_c = sheet[cell_c].value
            value_i = sheet[cell_i].value

            if (value_c is None or str(value_c).strip() == '') and (value_i is None or str(value_i).strip() == ''):
                empty_streak += 1
                if empty_streak >= 2:
                    break
            else:
                empty_streak = 0
                data['C'].append(value_c)
                data['I'].append(value_i)
                data['others'].append('請輸入文字')
            row += 1

    elif sheet_name == '表8.不確定分析':
        data = {
            'B': [],
            'C': [],
            'D': [],
            'E': [],
            'F': [],
            'G': [],
            'H': [],
            'I': [],
            'J': [],
            'K': [],
            'L': [],
            'M': [],
        }
        row = 4
        empty_streak = 0
        while True:
            cells = {col: sheet[f'{col}{row}'].value for col in list('BCDEFGHIJKLM')}
            if all((v is None or str(v).strip() == '') for v in cells.values()):
                empty_streak += 1
                if empty_streak >= 2:
                    break
            else:
                empty_streak = 0
                for col in data.keys():
                    data[col].append(cells.get(col))
            row += 1
    else:
        data = {}

    workbook.close()
    return data


def read_excel_data_pandas(excel_path, sheet_name):
    df = pd.read_excel(excel_path, sheet_name=sheet_name, header=2)
    data = {}
    if sheet_name == '表5.排放係數':
        df = df.dropna(subset=["排放類別"], how='all')
        gases = ["CO2", "CH4", "N2O", "HFCS", "PFCS", "SF6", "NF3"]
        transformed_rows = []
        for _, row in df.iterrows():
            has_valid_gas = any(pd.notna(row.get(gas)) and str(row.get(gas)).strip() != '' for gas in gases)
            if not has_valid_gas:
                continue
            for gas in gases:
                value = row.get(gas)
                if pd.isna(value) or str(value).strip() == '':
                    continue
                num = pd.to_numeric(value, errors='coerce')
                if pd.notna(num):
                    formatted_value = f"{num:.10f}"
                else:
                    formatted_value = str(value)
                transformed_rows.append({
                    "範疇或類別": row.get("排放類別", ""),
                    "排放源": row.get("排放源", ""),
                    "係數來源": row.get("係數來源", ""),
                    "係數名稱": row.get("係數名稱", ""),
                    "氣體": gas,
                    "溫室氣體排放係數": formatted_value,
                    "單位": row.get("單位", "")
                })
        final_df = pd.DataFrame(transformed_rows)
        final_df = final_df.fillna("")
        data = {
            '範疇或類別': final_df.get('範疇或類別', pd.Series([], dtype=str)).tolist(),
            '排放源': final_df.get('排放源', pd.Series([], dtype=str)).tolist(),
            '係數來源': final_df.get('係數來源', pd.Series([], dtype=str)).tolist(),
            '係數名稱': final_df.get('係數名稱', pd.Series([], dtype=str)).tolist(),
            '氣體': final_df.get('氣體', pd.Series([], dtype=str)).tolist(),
            '溫室氣體排放係數': final_df.get('溫室氣體排放係數', pd.Series([], dtype=str)).tolist(),
            '單位': final_df.get('單位', pd.Series([], dtype=str)).tolist()
        }
    return data


def read_excel_cell(excel_path, sheet_name, cell):
    try:
        workbook = load_workbook(excel_path, data_only=True)
        if sheet_name not in workbook.sheetnames:
            workbook.close()
            raise ValueError(f"Sheet '{sheet_name}' not found. Available: {workbook.sheetnames}")
        sheet = workbook[sheet_name]
        value = sheet[cell].value
        workbook.close()
        return str(value) if value is not None else ''
    except Exception as e:
        print(f"讀取儲存格 {cell} 失敗: {str(e)}")
        return ''


def read_excel_cells(excel_path, sheet_name, cells):
    try:
        workbook = load_workbook(excel_path, read_only=False, data_only=True)
        if sheet_name not in workbook.sheetnames:
            workbook.close()
            raise ValueError(f"Sheet '{sheet_name}' not found. Available: {workbook.sheetnames}")
        sheet = workbook[sheet_name]
        values = {cell: format_value(sheet[cell]) for cell in cells}
        workbook.close()
        return values
    except Exception as e:
        print(f"批量讀取儲存格失敗: {str(e)}")
        return {cell: '' for cell in cells}


def add_table_row(table):
    tr = OxmlElement('w:tr')
    for _ in range(len(table.columns)):
        tc = OxmlElement('w:tc')
        tc.append(OxmlElement('w:p'))
        tr.append(tc)
    table._tbl.append(tr)


def fill_word_table(word_path, output_path, table_index, excel_data, cell_mapping, start_row=0):
    doc = Document(word_path)
    if table_index >= len(doc.tables):
        raise IndexError(f"Template has only {len(doc.tables)} tables; requested index {table_index}")

    table = doc.tables[table_index]

    table.autofit = False
    table.allow_autofit = False

    # Set widths for all columns; extend defaults if needed
    num_cols = len(table.columns)
    column_widths = [COLUMN_WIDTH_DXA_DEFAULT] * num_cols
    for col_idx, width in enumerate(column_widths):
        for row in table.rows:
            cell = row.cells[col_idx]
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            # Remove existing tcW to avoid duplicates
            for child in list(tcPr):
                if child.tag == qn('w:tcW'):
                    tcPr.remove(child)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(width))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    max_data_len = max((len(excel_data.get(key, [])) for key in cell_mapping.keys()), default=0)

    required_rows = start_row + max_data_len
    while len(table.rows) < required_rows:
        add_table_row(table)

    for key, (row_offset, col) in cell_mapping.items():
        for i, value in enumerate(excel_data.get(key, [])):
            cell = table.cell(start_row + i, col)

            # Remove all existing paragraphs
            for para in list(cell.paragraphs):
                p = para._element
                p.getparent().remove(p)

            # Add new clean paragraph and run
            new_para = cell.add_paragraph()
            run = new_para.add_run(str(value).strip() if value is not None else '')
            _set_run_style(run)

            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            no_wrap = tcPr.find(qn('w:noWrap'))
            if no_wrap is not None:
                tcPr.remove(no_wrap)

    doc.save(output_path)


def replace_texts_in_word(word_path, output_path, replacements):
    doc = Document(word_path)

    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        if not original_text:
            continue
        new_text = original_text
        modified = False
        for old_text, new_text_value in replacements:
            if old_text in new_text:
                new_text = new_text.replace(old_text, new_text_value)
                modified = True
        if modified:
            _replace_paragraph_text(paragraph, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                if not original_text:
                    continue
                new_text = original_text
                modified = False
                for old_text, new_value in replacements:
                    if old_text in new_text:
                        new_text = new_text.replace(old_text, new_value)